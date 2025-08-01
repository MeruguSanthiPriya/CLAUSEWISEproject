import os
from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from typing import List, Set, Dict, Optional
import shutil
import requests
from bs4 import BeautifulSoup
import pdfplumber
import io
from langchain.schema import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
import re
import docx
import spacy
from datetime import datetime
import json
# New HF imports
from hf_model_manager import get_model_manager
from config import get_model_config, has_hf_token
# Google search imports
from googlesearch import search

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

PDF_DIR = "pdfs"
os.makedirs(PDF_DIR, exist_ok=True)

# --- Config ---
TRUSTED_DOMAINS = [".gov.in", ".nic.in", ".org", ".ac.in"]
MAX_PDFS = 10
MAX_HTMLS = 10
MAX_CHARS = 20000
CHUNK_SIZE = 4000  # For LLM context window

# --- Helper functions ---
def load_all_pdfs() -> List:
    docs = []
    if os.path.exists(PDF_DIR):
        for filename in os.listdir(PDF_DIR):
            if filename.endswith('.pdf'):
                loader = PDFPlumberLoader(os.path.join(PDF_DIR, filename))
                docs.extend(loader.load())
    return docs

def find_pdf_links(url):
    try:
        response = requests.get(url)
        response.raise_for_status()
    except Exception as e:
        return []
    soup = BeautifulSoup(response.text, "html.parser")
    pdf_links = []
    for link in soup.find_all("a", href=True):
        href = link["href"]
        if href.endswith(".pdf"):
            if href.startswith("http"):
                pdf_links.append(href)
            else:
                from urllib.parse import urljoin
                pdf_links.append(urljoin(url, href))
    return pdf_links

def download_pdf(pdf_url, save_dir):
    local_filename = os.path.join(save_dir, pdf_url.split("/")[-1])
    try:
        response = requests.get(pdf_url)
        response.raise_for_status()
        with open(local_filename, "wb") as f:
            f.write(response.content)
    except Exception:
        return None
    return local_filename

# --- Initialize HF Model Manager ---
model_manager = get_model_manager()

# --- Embeddings and Vectorstore ---
docs = load_all_pdfs()

# Initialize vectorstore and QA chain
vectorstore = None
retriever = None
qa_chain = None

if docs:
    try:
        # For API-based embeddings, we'll create a simple text-based retriever
        # since we can't use FAISS with API embeddings efficiently
        texts = [doc.page_content for doc in docs]
        metadatas = [doc.metadata for doc in docs]
        
        # Create a simple retriever that stores texts in memory
        class SimpleRetriever:
            def __init__(self, texts, metadatas):
                self.texts = texts
                self.metadatas = metadatas
            
            def get_relevant_documents(self, query: str, k: int = 4):
                # Simple keyword-based retrieval
                query_lower = query.lower()
                relevant_docs = []
                
                for i, text in enumerate(self.texts):
                    if any(word in text.lower() for word in query_lower.split()):
                        relevant_docs.append(Document(
                            page_content=text,
                            metadata=metadatas[i]
                        ))
                
                return relevant_docs[:k]
        
        retriever = SimpleRetriever(texts, metadatas)
        
        # Get primary LLM for QA
        llm = model_manager.get_llm("primary")
        if llm:
            qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
            print("✅ QA chain initialized successfully")
        else:
            print("⚠️ No LLM available for QA chain")
    except Exception as e:
        print(f"❌ Error initializing retriever/QA chain: {e}")
else:
    print("⚠️ No documents available")

class Question(BaseModel):
    question: str

# --- Google Search Helper ---
def google_search_links(query: str, num_results: int = 10) -> List[str]:
    try:
        from googlesearch import search
    except ImportError:
        raise RuntimeError("Please install googlesearch-python: pip install googlesearch-python")
    return list(search(query, num_results=num_results))

def is_trusted_domain(url: str) -> bool:
    return any(domain in url for domain in TRUSTED_DOMAINS)

# --- Scrape HTML for PDF links ---
def extract_pdf_links_from_html(url: str) -> Set[str]:
    pdf_links = set()
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for a in soup.find_all("a", href=True):
            href = a["href"]
            if href.lower().endswith(".pdf"):
                if href.startswith("http"):
                    pdf_links.add(href)
                else:
                    from urllib.parse import urljoin
                    pdf_links.add(urljoin(url, href))
    except Exception:
        pass
    return pdf_links

# --- PDF Text Extraction Helper ---
def extract_text_from_pdf_url(pdf_url: str) -> Dict:
    try:
        response = requests.get(pdf_url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        response.raise_for_status()
        with pdfplumber.open(io.BytesIO(response.content)) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
        return {"content": text, "source": pdf_url}
    except Exception:
        return {"content": "", "source": pdf_url}

# --- HTML Text Extraction Helper ---
def extract_text_from_html_url(url: str) -> Dict:
    try:
        resp = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "header", "footer", "nav"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)
        return {"content": text, "source": url}
    except Exception:
        return {"content": "", "source": url}

# --- Chunking Helper ---
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> List[str]:
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

# --- NEW: Document Analysis Functions ---

def extract_text_from_content(content: bytes, filename: str) -> str:
    """Extract text from uploaded file content"""
    try:
        if filename.lower().endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                return text
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        elif filename.lower().endswith('.txt'):
            return content.decode('utf-8')
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text: {str(e)}")

def extract_clauses(text: str) -> List[Dict]:
    """Enhanced clause extraction and breakdown from legal documents"""
    clauses = []
    
    # Comprehensive clause detection patterns
    clause_patterns = [
        # Obligation clauses
        r'(?:^|\n)([A-Z][^.]*?(?:shall|will|must|agrees?|acknowledges?|represents?|warrants?|obligated|required|responsible)[^.]*?\.)',
        
        # Party-related clauses
        r'(?:^|\n)([A-Z][^.]*?(?:party|parties|company|corporation|individual|employer|employee|tenant|landlord|buyer|seller)[^.]*?\.)',
        
        # Confidentiality clauses
        r'(?:^|\n)([A-Z][^.]*?(?:confidential|proprietary|intellectual property|trade secret|non-disclosure)[^.]*?\.)',
        
        # Termination clauses
        r'(?:^|\n)([A-Z][^.]*?(?:termination|breach|default|cancellation|expiration)[^.]*?\.)',
        
        # Payment clauses
        r'(?:^|\n)([A-Z][^.]*?(?:payment|compensation|salary|fee|amount|remuneration|bonus)[^.]*?\.)',
        
        # Liability clauses
        r'(?:^|\n)([A-Z][^.]*?(?:liability|damages|indemnification|warranty|guarantee)[^.]*?\.)',
        
        # Duration clauses
        r'(?:^|\n)([A-Z][^.]*?(?:term|duration|period|effective|commencement|expiry)[^.]*?\.)',
        
        # Jurisdiction clauses
        r'(?:^|\n)([A-Z][^.]*?(?:jurisdiction|governing law|venue|applicable law)[^.]*?\.)',
        
        # Amendment clauses
        r'(?:^|\n)([A-Z][^.]*?(?:amendment|modification|change|alteration)[^.]*?\.)',
        
        # Force majeure clauses
        r'(?:^|\n)([A-Z][^.]*?(?:force majeure|act of god|unforeseen|circumstances)[^.]*?\.)',
        
        # Assignment clauses
        r'(?:^|\n)([A-Z][^.]*?(?:assignment|transfer|delegation|subcontract)[^.]*?\.)',
        
        # Severability clauses
        r'(?:^|\n)([A-Z][^.]*?(?:severability|severable|invalid|unenforceable)[^.]*?\.)',
        
        # Entire agreement clauses
        r'(?:^|\n)([A-Z][^.]*?(?:entire agreement|complete|supersede|integrated)[^.]*?\.)',
        
        # Waiver clauses
        r'(?:^|\n)([A-Z][^.]*?(?:waiver|waive|forfeit|relinquish)[^.]*?\.)',
        
        # Notice clauses
        r'(?:^|\n)([A-Z][^.]*?(?:notice|notification|inform|advise)[^.]*?\.)',
        
        # Arbitration clauses
        r'(?:^|\n)([A-Z][^.]*?(?:arbitration|dispute|mediation|resolution)[^.]*?\.)',
    ]
    
    for pattern in clause_patterns:
        matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
        for match in matches:
            clause_text = match.group(1).strip()
            if len(clause_text) > 20:  # Filter out very short matches
                clauses.append({
                    "original": clause_text,
                    "type": classify_clause_type(clause_text),
                    "importance": assess_clause_importance(clause_text)
                })
    
    # Remove duplicates while preserving order
    seen = set()
    unique_clauses = []
    for clause in clauses:
        if clause["original"] not in seen:
            seen.add(clause["original"])
            unique_clauses.append(clause)
    
    # Sort by importance
    unique_clauses.sort(key=lambda x: x["importance"], reverse=True)
    
    return unique_clauses[:25]  # Increased limit for better coverage

def classify_clause_type(clause_text: str) -> str:
    """Enhanced classification of legal clause types"""
    clause_lower = clause_text.lower()
    
    if any(word in clause_lower for word in ['confidential', 'proprietary', 'trade secret', 'non-disclosure']):
        return "Confidentiality"
    elif any(word in clause_lower for word in ['payment', 'compensation', 'salary', 'fee', 'amount', 'remuneration']):
        return "Payment"
    elif any(word in clause_lower for word in ['termination', 'breach', 'default', 'cancellation']):
        return "Termination"
    elif any(word in clause_lower for word in ['liability', 'damages', 'indemnification', 'warranty']):
        return "Liability"
    elif any(word in clause_lower for word in ['intellectual property', 'patent', 'copyright', 'trademark']):
        return "Intellectual Property"
    elif any(word in clause_lower for word in ['party', 'parties', 'company', 'corporation', 'employer', 'employee']):
        return "Parties"
    elif any(word in clause_lower for word in ['shall', 'will', 'must', 'agrees', 'acknowledges', 'obligated']):
        return "Obligation"
    elif any(word in clause_lower for word in ['jurisdiction', 'governing law', 'venue']):
        return "Jurisdiction"
    elif any(word in clause_lower for word in ['amendment', 'modification', 'change']):
        return "Amendment"
    elif any(word in clause_lower for word in ['force majeure', 'act of god', 'unforeseen']):
        return "Force Majeure"
    elif any(word in clause_lower for word in ['assignment', 'transfer', 'delegation']):
        return "Assignment"
    elif any(word in clause_lower for word in ['severability', 'severable', 'invalid']):
        return "Severability"
    elif any(word in clause_lower for word in ['entire agreement', 'complete', 'supersede']):
        return "Entire Agreement"
    elif any(word in clause_lower for word in ['waiver', 'waive', 'forfeit']):
        return "Waiver"
    elif any(word in clause_lower for word in ['notice', 'notification', 'inform']):
        return "Notice"
    elif any(word in clause_lower for word in ['arbitration', 'dispute', 'mediation']):
        return "Dispute Resolution"
    elif any(word in clause_lower for word in ['term', 'duration', 'period', 'effective']):
        return "Duration"
    else:
        return "General"

def assess_clause_importance(clause_text: str) -> int:
    """Assess the importance of a legal clause (higher score = more important)"""
    clause_lower = clause_text.lower()
    importance_score = 0
    
    # High importance keywords
    high_importance = ['shall', 'must', 'liability', 'damages', 'termination', 'breach', 'confidential']
    for word in high_importance:
        if word in clause_lower:
            importance_score += 3
    
    # Medium importance keywords
    medium_importance = ['agrees', 'acknowledges', 'payment', 'compensation', 'jurisdiction', 'governing law']
    for word in medium_importance:
        if word in clause_lower:
            importance_score += 2
    
    # Low importance keywords
    low_importance = ['notice', 'amendment', 'waiver', 'severability', 'entire agreement']
    for word in low_importance:
        if word in clause_lower:
            importance_score += 1
    
    # Bonus for longer clauses (more detailed)
    if len(clause_text) > 100:
        importance_score += 1
    
    return importance_score

def simplify_clause_with_llm(clause_text: str) -> str:
    """Simplify legal clause using LLM with better formatting"""
    try:
        # Create a better prompt for clause simplification with enhanced formatting
        prompt = f"""
        Simplify this legal clause into plain English with excellent formatting:
        
        Original: {clause_text}
        
        Provide a clear, well-formatted explanation with:
        - Clear headings
        - Bullet points for key elements
        - Bold text for important terms
        - Proper spacing and structure
        - Easy-to-understand language
        """
        
        # Use the model manager to get LLM and generate response
        response = model_manager.generate_text(prompt, model_type="primary")
        
        # If the response doesn't have good formatting, enhance it
        if response and not any(marker in response.lower() for marker in ['•', 'bullet', '-', '*', '**']):
            # Add bullet point formatting
            lines = response.split('.')
            formatted_lines = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 10:  # Only format substantial lines
                    formatted_lines.append(f"• **{line}**")
            
            if formatted_lines:
                return "**📋 Clause Analysis:**\n\n" + "\n".join(formatted_lines)
        
        return response if response else basic_clause_simplification(clause_text)
    except Exception as e:
        print(f"Error in clause simplification: {e}")
        return basic_clause_simplification(clause_text)

def basic_clause_simplification(clause_text: str) -> str:
    """Basic clause simplification without LLM with professional formatting"""
    # Simple replacements for common legal terms
    replacements = {
        'shall': 'must',
        'hereby': '',
        'thereof': 'of this',
        'therein': 'in this',
        'thereto': 'to this',
        'whereas': 'since',
        'pursuant to': 'according to',
        'in accordance with': 'following',
        'notwithstanding': 'despite',
        'provided that': 'as long as',
        'subject to': 'depending on',
        'in the event that': 'if',
        'prior to': 'before',
        'subsequent to': 'after',
        'terminate': 'end',
        'breach': 'violation',
        'default': 'failure to comply',
        'liability': 'responsibility',
        'damages': 'compensation',
        'indemnification': 'protection from loss',
        'confidential': 'private',
        'proprietary': 'owned by the company',
        'intellectual property': 'ideas and creations',
    }
    
    simplified = clause_text
    for legal_term, plain_term in replacements.items():
        simplified = re.sub(r'\b' + legal_term + r'\b', plain_term, simplified, flags=re.IGNORECASE)
    
    # Clean up the text
    simplified = re.sub(r'\s+', ' ', simplified)  # Remove extra spaces
    simplified = simplified.strip()
    
    # Extract key information
    key_info = []
    
    # Enhanced amount extraction - multiple patterns to catch all formats
    amount_patterns = [
        r'(?:rs\.?|₹|rupees?|inr)\s*[:\-]?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)',  # Rs. 50,000
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rs\.?|₹|rupees?|inr)',  # 50,000 Rs.
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rupees?|rs?|inr)',  # 50000 rupees
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:lakh|lacs)',  # 5 lakhs
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:crore|crores)',  # 1 crore
        r'\b(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rs?|rupees?|inr)\b',  # 15000 rs
        r'\b(?:rs?|rupees?|inr)\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',  # rs 15000
        r'\b(\d+(?:,\d{3})*(?:\.\d{2})?)\b',  # Any number with commas (fallback)
    ]
    
    amounts = []
    for pattern in amount_patterns:
        found = re.findall(pattern, clause_text, re.IGNORECASE)
        if found:
            amounts.extend(found)
            break
    
    if amounts:
        key_info.append(f"**💰 Amount:** ₹{amounts[0]}")
    else:
        # Try to find any number that might be an amount
        number_pattern = r'\b(\d+(?:,\d{3})*(?:\.\d{2})?)\b'
        numbers = re.findall(number_pattern, clause_text)
        if numbers:
            key_info.append(f"**💰 Amount:** ₹{numbers[0]}")
    
    # Extract dates
    dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}\b|\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b', simplified, re.IGNORECASE)
    if dates:
        key_info.append(f"**📅 Date:** {dates[0]}")
    
    # Extract parties
    parties = re.findall(r'\b(?:party|parties|company|corporation|individual|employer|employee|landlord|tenant|buyer|seller|lessor|lessee)\b', simplified, re.IGNORECASE)
    if parties:
        key_info.append(f"**👥 Parties:** {parties[0].title()}")
    
    # Extract obligations
    obligations = re.findall(r'\b(?:shall|will|must|agrees?|acknowledges?|represents?|warrants?|obligated|responsible)\b', simplified, re.IGNORECASE)
    if obligations:
        key_info.append(f"**📋 Obligation:** {obligations[0].title()}")
    
    if key_info:
        return "**🔍 Simplified Analysis:**\n\n" + "\n".join(key_info)
    
    # If no key info found, return a simple summary
    sentences = simplified.split('.')
    if sentences and len(sentences[0]) > 10:
        return f"**🔍 Simplified Analysis:**\n\n{sentences[0].strip()}"
    
    return f"**🔍 Simplified Analysis:**\n\n{simplified}"

def simplify_specific_clause(clause_text: str) -> str:
    """Simplify a specific clause from the document with comprehensive analysis"""
    try:
        # Extract key information from the specific clause
        clause_lower = clause_text.lower()
        
        # Enhanced amount extraction - multiple patterns to catch all formats
        amount_patterns = [
            r'(?:rs\.?|₹|rupees?|inr)\s*[:\-]?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)',  # Rs. 50,000
            r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rs\.?|₹|rupees?|inr)',  # 50,000 Rs.
            r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rupees?|rs?|inr)',  # 50000 rupees
            r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:lakh|lacs)',  # 5 lakhs
            r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:crore|crores)',  # 1 crore
        ]
        
        amounts = []
        for pattern in amount_patterns:
            found = re.findall(pattern, clause_text, re.IGNORECASE)
            if found:
                amounts.extend(found)
                break
        
        # Enhanced clause analysis with detailed explanations
        if "commence" in clause_lower or "start" in clause_lower or "begin" in clause_lower or "effective" in clause_lower:
            # Duration/commencement clause
            dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}\b|\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b', clause_text, re.IGNORECASE)
            durations = re.findall(r'\b\d+\s+(?:month|year|day|week)s?\b', clause_lower)
            
            simplified = "**📅 Duration & Commencement Clause**\n\n"
            simplified += "**Key Details:**\n"
            if dates:
                simplified += f"• **Effective Date:** {dates[0]}\n"
                if len(dates) > 1:
                    simplified += f"• **End Date:** {dates[1]}\n"
            if durations:
                simplified += f"• **Duration Period:** {durations[0]}\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• This clause establishes the timeline for the agreement\n"
            simplified += "• Defines when obligations begin and end\n"
            simplified += "• Important for calculating deadlines and compliance\n"
            
        elif "rent" in clause_lower or "lease" in clause_lower or "tenancy" in clause_lower:
            # Rental/lease clause
            parties = re.findall(r'\b(?:landlord|tenant|lessor|lessee|owner|renter)\b', clause_lower)
            addresses = re.findall(r'\d+\s+[A-Za-z\s]+(?:street|road|avenue|lane|drive|residency|apartment|flat|building)', clause_text, re.IGNORECASE)
            
            simplified = "**🏠 Rental/Lease Agreement Clause**\n\n"
            simplified += "**Key Details:**\n"
            if parties:
                simplified += f"• **Involved Parties:** {' and '.join(parties)}\n"
            if addresses:
                simplified += f"• **Property Address:** {addresses[0]}\n"
            if amounts:
                simplified += f"• **Rental Amount:** ₹{amounts[0]}\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Establishes landlord-tenant relationship\n"
            simplified += "• Defines rent payment obligations\n"
            simplified += "• Governs property usage rights\n"
            simplified += "• Subject to Rent Control Act provisions\n"
                
        elif "payment" in clause_lower or "pay" in clause_lower or "compensation" in clause_lower:
            # Payment clause
            dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2},?\s+\d{4}\b', clause_text, re.IGNORECASE)
            payment_terms = re.findall(r'(?:within|before|after|on|by)\s+\d+\s+(?:day|week|month)s?', clause_lower)
            
            simplified = "**💰 Payment Terms Clause**\n\n"
            simplified += "**Key Details:**\n"
            if amounts:
                simplified += f"• **Payment Amount:** ₹{amounts[0]}\n"
            if dates:
                simplified += f"• **Due Date:** {dates[0]}\n"
            if payment_terms:
                simplified += f"• **Payment Timeline:** {payment_terms[0]}\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Creates enforceable payment obligations\n"
            simplified += "• Late payment may incur penalties\n"
            simplified += "• Subject to Indian Contract Act, 1872\n"
            simplified += "• May include interest on delayed payments\n"
                
        elif "terminate" in clause_lower or "end" in clause_lower or "cancel" in clause_lower or "breach" in clause_lower:
            # Termination clause
            conditions = re.findall(r'(?:if|when|upon|in case of|in the event of)\s+[^.]*', clause_lower)
            notice_periods = re.findall(r'\d+\s+(?:day|week|month)s?\s+(?:notice|advance)', clause_lower)
            
            simplified = "**🚪 Termination & Breach Clause**\n\n"
            simplified += "**Key Details:**\n"
            if conditions:
                simplified += f"• **Termination Conditions:** {conditions[0]}\n"
            if notice_periods:
                simplified += f"• **Notice Period:** {notice_periods[0]}\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Defines grounds for contract termination\n"
            simplified += "• Establishes breach consequences\n"
            simplified += "• May require notice periods\n"
            simplified += "• Subject to specific performance remedies\n"
                
        elif "confidential" in clause_lower or "secret" in clause_lower or "proprietary" in clause_lower:
            # Confidentiality clause
            simplified = "**🔒 Confidentiality & Non-Disclosure Clause**\n\n"
            simplified += "**Key Details:**\n"
            simplified += "• **Scope:** Covers proprietary and confidential information\n"
            simplified += "• **Duration:** Typically extends beyond agreement termination\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Creates legal obligation to maintain secrecy\n"
            simplified += "• Breach may result in injunctive relief\n"
            simplified += "• Subject to intellectual property laws\n"
            simplified += "• May include liquidated damages\n"
            
        elif "liability" in clause_lower or "damage" in clause_lower or "indemnification" in clause_lower:
            # Liability clause
            simplified = "**⚖️ Liability & Indemnification Clause**\n\n"
            simplified += "**Key Details:**\n"
            simplified += "• **Scope:** Defines responsibility for damages\n"
            simplified += "• **Limitations:** May cap liability amounts\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Allocates risk between parties\n"
            simplified += "• May limit or exclude certain damages\n"
            simplified += "• Subject to reasonableness tests\n"
            simplified += "• Important for insurance considerations\n"
            
        elif "jurisdiction" in clause_lower or "governing law" in clause_lower or "venue" in clause_lower:
            # Jurisdiction clause
            simplified = "**🏛️ Jurisdiction & Governing Law Clause**\n\n"
            simplified += "**Key Details:**\n"
            simplified += "• **Applicable Law:** Determines legal framework\n"
            simplified += "• **Dispute Resolution:** Specifies court jurisdiction\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Determines which laws apply\n"
            simplified += "• Establishes dispute resolution forum\n"
            simplified += "• Critical for enforcement\n"
            simplified += "• May affect choice of law\n"
            
        else:
            # General clause - extract key terms and amounts
            key_terms = re.findall(r'\b(?:shall|will|must|agrees?|acknowledges?|represents?|warrants?|obligated|responsible)\b', clause_lower)
            parties = re.findall(r'\b(?:party|parties|company|corporation|individual|employer|employee|buyer|seller)\b', clause_lower)
            
            simplified = "**📋 General Legal Clause**\n\n"
            simplified += "**Key Details:**\n"
            if key_terms:
                simplified += f"• **Primary Obligation:** {key_terms[0]}\n"
            if parties:
                simplified += f"• **Involved Parties:** {parties[0]}\n"
            if amounts:
                simplified += f"• **Monetary Amount:** ₹{amounts[0]}\n"
            simplified += "\n**Legal Implications:**\n"
            simplified += "• Creates binding legal obligations\n"
            simplified += "• Subject to contract law principles\n"
            simplified += "• May have enforcement mechanisms\n"
            simplified += "• Important for compliance tracking\n"
        
        return simplified
        
    except Exception as e:
        # Fallback to basic simplification
        return basic_clause_simplification(clause_text)

def extract_entities(text: str) -> Dict:
    """Enhanced Named Entity Recognition (NER) for legal documents"""
    entities = {
        "parties": [],
        "dates": [],
        "amounts": [],
        "locations": [],
        "legal_terms": [],
        "obligations": [],
        "penalties": [],
        "jurisdictions": []
    }
    
    # Enhanced party extraction
    party_patterns = [
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc\.|Corp\.|LLC|Ltd\.|Company|Corporation|Limited|Partnership|Associates)\b',
        r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b',  # Simple name pattern
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:LLP|PLC|PLLC)\b',
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Group|Enterprises|Solutions|Technologies)\b',
    ]
    
    for pattern in party_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            party = match.group(0)
            if party not in entities["parties"] and len(party) > 3:
                entities["parties"].append(party)
    
    # Enhanced date extraction
    date_patterns = [
        r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
        r'\b\d{4}-\d{2}-\d{2}\b',
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},?\s+\d{4}\b',
        r'\b(?:effective|commencing|starting)\s+(?:on|from|as\s+of)\s+\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b',
    ]
    
    for pattern in date_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            date = match.group(0)
            if date not in entities["dates"]:
                entities["dates"].append(date)
    
    # Enhanced monetary amount extraction with comprehensive patterns
    amount_patterns = [
        # US Dollar patterns
        r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?',
        r'\$\d+(?:\.\d{2})?',
        r'\b\d+(?:\.\d{2})?\s*(?:dollars?|USD)\b',
        
        # Indian Rupee patterns - comprehensive coverage
        r'(?:rs\.?|₹|rupees?|inr)\s*[:\-]?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)',  # Rs. 50,000
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rs\.?|₹|rupees?|inr)',  # 50,000 Rs.
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rupees?|rs?|inr)',  # 50000 rupees
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:lakh|lacs)',  # 5 lakhs
        r'(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:crore|crores)',  # 1 crore
        r'Rs\.\s*\d{1,3}(?:,\d{3})*(?:\.\d{2})?',
        r'Rs\.\s*\d+(?:\.\d{2})?',
        r'\b\d+(?:\.\d{2})?\s*(?:rupees?|INR|Rs\.?)\b',
        
        # Contextual amount patterns
        r'\b(?:amount|sum|payment|compensation|salary|fee|rent|lease|deposit)\s+of\s+(?:rs\.?|₹|rupees?|inr)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',
        r'\b(?:annual|monthly|weekly|daily|yearly)\s+(?:salary|compensation|payment|rent|fee)\s+of\s+(?:rs\.?|₹|rupees?|inr)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',
        r'\b(?:total|aggregate|maximum|minimum)\s+(?:amount|sum|payment)\s+of\s+(?:rs\.?|₹|rupees?|inr)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',
        
        # Penalty and fine patterns
        r'\b(?:penalty|fine|damages|compensation)\s+of\s+(?:rs\.?|₹|rupees?|inr)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',
        r'\b(?:liquidated|punitive|exemplary)\s+(?:damages|compensation)\s+of\s+(?:rs\.?|₹|rupees?|inr)?\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',
        
        # Additional patterns for better coverage
        r'\b(\d+(?:,\d{3})*(?:\.\d{2})?)\s*(?:rs?|rupees?|inr)\b',  # 15000 rs
        r'\b(?:rs?|rupees?|inr)\s*(\d+(?:,\d{3})*(?:\.\d{2})?)\b',  # rs 15000
        r'\b(\d+(?:,\d{3})*(?:\.\d{2})?)\b',  # Any number with commas (fallback)
    ]
    
    for pattern in amount_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            amount = match.group(0)
            # Clean up the amount for better display
            if amount not in entities["amounts"]:
                # Format amounts consistently
                if 'rs' in amount.lower() or 'rupee' in amount.lower() or '₹' in amount or 'inr' in amount.lower():
                    # Extract the numeric part and format as ₹
                    numeric_part = re.search(r'(\d+(?:,\d{3})*(?:\.\d{2})?)', amount)
                    if numeric_part:
                        formatted_amount = f"₹{numeric_part.group(1)}"
                        if formatted_amount not in entities["amounts"]:
                            entities["amounts"].append(formatted_amount)
                elif any(word in amount.lower() for word in ['dollar', 'usd', '$']):
                    # Format USD amounts
                    numeric_part = re.search(r'(\d+(?:,\d{3})*(?:\.\d{2})?)', amount)
                    if numeric_part:
                        formatted_amount = f"${numeric_part.group(1)}"
                        if formatted_amount not in entities["amounts"]:
                            entities["amounts"].append(formatted_amount)
                else:
                    # For any other numeric amounts, ensure they're properly formatted
                    numeric_part = re.search(r'(\d+(?:,\d{3})*(?:\.\d{2})?)', amount)
                    if numeric_part:
                        formatted_amount = f"₹{numeric_part.group(1)}"
                        if formatted_amount not in entities["amounts"]:
                            entities["amounts"].append(formatted_amount)
                    else:
                        entities["amounts"].append(amount)
    
    # Enhanced location extraction
    location_patterns = [
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s+[A-Z]{2}\b',  # City, State
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',  # City, Country
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:County|District|Province|State)\b',
        r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Street|Avenue|Road|Boulevard|Lane)\b',
    ]
    
    for pattern in location_patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            location = match.group(0)
            if location not in entities["locations"]:
                entities["locations"].append(location)
    
    # Enhanced legal terms extraction
    legal_terms = [
        'confidentiality', 'non-disclosure', 'intellectual property', 'liability',
        'indemnification', 'breach', 'termination', 'default', 'damages',
        'arbitration', 'governing law', 'jurisdiction', 'force majeure',
        'non-compete', 'non-solicitation', 'severability', 'entire agreement',
        'amendment', 'waiver', 'assignment', 'subrogation', 'estoppel',
        'rescission', 'specific performance', 'liquidated damages', 'penalty',
        'good faith', 'reasonable care', 'due diligence', 'material breach'
    ]
    
    for term in legal_terms:
        if re.search(r'\b' + term + r'\b', text, re.IGNORECASE):
            if term not in entities["legal_terms"]:
                entities["legal_terms"].append(term)
    
    # Extract obligations
    obligation_patterns = [
        r'\b(?:shall|will|must|agrees?|acknowledges?|represents?|warrants?)\s+[^.]*?\.',
        r'\b(?:obligated|required|responsible|duty|obligation)\s+[^.]*?\.',
    ]
    
    for pattern in obligation_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            obligation = match.group(0)
            if obligation not in entities["obligations"]:
                entities["obligations"].append(obligation)
    
    # Extract penalties
    penalty_patterns = [
        r'\b(?:penalty|fine|sanction|consequence)\s+[^.]*?\.',
        r'\b(?:liquidated\s+damages|punitive\s+damages)\s+[^.]*?\.',
    ]
    
    for pattern in penalty_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            penalty = match.group(0)
            if penalty not in entities["penalties"]:
                entities["penalties"].append(penalty)
    
    # Extract jurisdictions
    jurisdiction_patterns = [
        r'\b(?:jurisdiction|venue|governing\s+law)\s+[^.]*?\.',
        r'\b(?:laws?\s+of\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\b',
    ]
    
    for pattern in jurisdiction_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            jurisdiction = match.group(0)
            if jurisdiction not in entities["jurisdictions"]:
                entities["jurisdictions"].append(jurisdiction)
    
    return entities

def classify_document(text: str) -> str:
    """Classify document type"""
    text_lower = text.lower()
    
    # Document type classification based on keywords
    doc_types = {
        "Non-Disclosure Agreement (NDA)": [
            'confidential', 'non-disclosure', 'trade secret', 'proprietary information',
            'disclose', 'confidentiality agreement'
        ],
        "Employment Contract": [
            'employment', 'employee', 'employer', 'salary', 'compensation',
            'job duties', 'work schedule', 'benefits'
        ],
        "Service Agreement": [
            'service', 'provider', 'client', 'scope of work', 'deliverables',
            'service level', 'performance'
        ],
        "Lease Agreement": [
            'lease', 'tenant', 'landlord', 'rent', 'property', 'premises',
            'rental agreement', 'leasehold'
        ],
        "Purchase Agreement": [
            'purchase', 'buyer', 'seller', 'purchase price', 'closing date',
            'title', 'escrow'
        ],
        "Partnership Agreement": [
            'partnership', 'partner', 'joint venture', 'profit sharing',
            'partnership interest'
        ],
        "License Agreement": [
            'license', 'licensor', 'licensee', 'intellectual property',
            'royalty', 'permitted use'
        ]
    }
    
    scores = {}
    for doc_type, keywords in doc_types.items():
        score = sum(1 for keyword in keywords if keyword in text_lower)
        scores[doc_type] = score
    
    # Return the document type with highest score, or "General Legal Document"
    if scores:
        best_type = max(scores, key=scores.get)
        if scores[best_type] > 0:
            return best_type
    
    return "General Legal Document"

def extract_and_simplify_clauses(text: str) -> List[Dict]:
    """Extract and simplify legal clauses"""
    clauses = extract_clauses(text)
    
    # Simplify each clause using our specific clause simplification
    for clause in clauses:
        clause["simplified"] = simplify_specific_clause(clause["original"])
    
    return clauses

@app.post("/ask")
async def ask_question(q: Question):
    """Answer legal questions using available documents and knowledge with comprehensive sources"""
    try:
        # Get question-specific sources using Google search
        relevant_sources = get_question_specific_sources(q.question)
        
        # Extract content from URLs for research-based answer
        research_content = ""
        if relevant_sources:
            research_content = extract_content_from_urls(relevant_sources)
        
        # Enhanced prompt for research-based legal answers
        if research_content.strip():
            prompt = f"""
            Answer this legal question directly: {q.question}
            
            Research Content:
            {research_content}
            
            Provide a clear, direct answer that:
            1. Gives a specific answer to the question asked
            2. Uses information from the research content when available
            3. Provides practical legal information based on Indian law
            4. Includes relevant legal citations and references
            5. Uses bullet points for key information
            
            Format your response with:
            - Direct answer to the question first
            - Relevant legal information and citations
            - Practical steps or procedures if applicable
            - Important deadlines or time limits if relevant
            - Bullet points for key points
            
            Always provide a helpful, informative answer based on Indian legal knowledge. Do not say you cannot answer or provide fallback responses.
            """
        else:
            prompt = f"""
            Answer this legal question directly and clearly: {q.question}
            
            Provide a clear, direct answer that:
            1. Gives a specific answer to the question asked, even if you have to make a best guess based on Indian law and common practice
            2. Uses layman language and is easy to understand
            3. Provides practical legal information based on Indian law
            4. Includes relevant legal citations and references if possible
            5. Uses bullet points for key information
            
            Format your response with:
            - Direct answer to the question first
            - Relevant legal information and citations
            - Practical steps or procedures if applicable
            - Important deadlines or time limits if relevant
            - Bullet points for key points
            
            Do not say you cannot answer. Always provide a helpful, informative answer based on Indian legal knowledge, even if it is a best guess.
            """
        
        # Get response directly from model manager
        answer = model_manager.generate_text(prompt, model_type="primary")
        
        # Get question-specific sources using Google search
        relevant_sources = get_question_specific_sources(q.question)
        
        # If no Google search results, fall back to document-based sources
        if not relevant_sources and docs:
            question_words = q.question.lower().split()
            question_keywords = [word for word in question_words if len(word) > 3]
            
            # Enhanced relevance scoring
            doc_scores = []
            for doc in docs:
                doc_content = doc.page_content.lower()
                relevance_score = 0
                
                # Score based on keyword matches
                for keyword in question_keywords:
                    if keyword in doc_content:
                        relevance_score += 2
                
                # Bonus for exact phrase matches
                if any(phrase in doc_content for phrase in question_words if len(phrase) > 5):
                    relevance_score += 3
                
                # Bonus for legal terms
                legal_terms = ['act', 'section', 'clause', 'article', 'regulation', 'statute', 'law']
                for term in legal_terms:
                    if term in doc_content and term in question.lower():
                        relevance_score += 2
                
                if relevance_score > 0:
                    doc_scores.append({
                        'source': doc.metadata.get("source", "Unknown"),
                        'score': relevance_score,
                        'content': doc.page_content[:200]  # First 200 chars for context
                    })
            
            # Sort by relevance score and take top sources
            doc_scores.sort(key=lambda x: x['score'], reverse=True)
            relevant_sources = [doc['source'] for doc in doc_scores[:5]]
        
        # If still no sources, add basic legal sources based on question type
        if not relevant_sources:
            question_lower = q.question.lower()
            
            # Determine question type and add relevant sources
            if any(word in question_lower for word in ['contract', 'agreement', 'breach']):
                relevant_sources.extend([
                    "Indian Contract Act, 1872",
                    "Specific Relief Act, 1963",
                    "Sale of Goods Act, 1930"
                ])
            elif any(word in question_lower for word in ['consumer', 'complaint', 'defect']):
                relevant_sources.extend([
                    "Consumer Protection Act, 2019",
                    "Consumer Protection Rules, 2020",
                    "National Consumer Disputes Redressal Commission"
                ])
            elif any(word in question_lower for word in ['employment', 'labour', 'worker', 'employee']):
                relevant_sources.extend([
                    "Industrial Disputes Act, 1947",
                    "Factories Act, 1948",
                    "Minimum Wages Act, 1948",
                    "Payment of Wages Act, 1936"
                ])
            elif any(word in question_lower for word in ['property', 'real estate', 'land', 'rent']):
                relevant_sources.extend([
                    "Transfer of Property Act, 1882",
                    "Registration Act, 1908",
                    "Rent Control Act (State-specific)",
                    "Real Estate (Regulation and Development) Act, 2016"
                ])
            elif any(word in question_lower for word in ['criminal', 'offence', 'penalty', 'punishment']):
                relevant_sources.extend([
                    "Indian Penal Code, 1860",
                    "Code of Criminal Procedure, 1973",
                    "Evidence Act, 1872"
                ])
            elif any(word in question_lower for word in ['family', 'marriage', 'divorce', 'maintenance']):
                relevant_sources.extend([
                    "Hindu Marriage Act, 1955",
                    "Special Marriage Act, 1954",
                    "Hindu Succession Act, 1956",
                    "Guardian and Wards Act, 1890"
                ])
            elif any(word in question_lower for word in ['company', 'corporate', 'business', 'incorporation']):
                relevant_sources.extend([
                    "Companies Act, 2013",
                    "Limited Liability Partnership Act, 2008",
                    "Partnership Act, 1932"
                ])
            else:
                # General legal sources
                relevant_sources.extend([
                    "Constitution of India",
                    "Code of Civil Procedure, 1908",
                    "Indian Evidence Act, 1872",
                    "Limitation Act, 1963"
                ])
        
        # Remove duplicates while preserving order
        seen = set()
        unique_sources = []
        for source in relevant_sources:
            if source not in seen:
                seen.add(source)
                unique_sources.append(source)
        
        # Limit to top 8 most relevant sources
        final_sources = unique_sources[:8]
        
        return {
            "answer": answer, 
            "sources": final_sources
        }
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error in ask_question: {e}")
        raise HTTPException(status_code=500, detail=f"Error answering question: {e}")



def get_question_specific_sources(question: str) -> List[str]:
    """Get question-specific sources using Google search"""
    try:
        # Create search queries based on question content
        search_terms = [
            f"{question} Indian law PDF",
            f"{question} legal document India",
            f"{question} government notification India",
            f"{question} legal guidelines India"
        ]
        
        sources = []
        
        # Search for each term
        for search_term in search_terms:
            try:
                # Search for PDFs and legal documents
                search_results = list(search(search_term + " filetype:pdf OR site:gov.in OR site:legislative.gov.in OR site:indiankanoon.org", num_results=5))
                sources.extend(search_results)
                
            except Exception as e:
                print(f"Search error for term '{search_term}': {e}")
                continue
        
        # Remove duplicates and limit results
        unique_sources = list(dict.fromkeys(sources))  # Preserve order
        return unique_sources[:10]  # Return top 10 sources
        
    except Exception as e:
        print(f"Error in get_question_specific_sources: {e}")
        return []

def extract_content_from_urls(urls: List[str]) -> str:
    """Extract content from URLs for research-based answers"""
    content = ""
    
    for url in urls[:5]:  # Limit to first 5 URLs
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text content
                text = soup.get_text()
                
                # Clean up text
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                content += f"\n\nSource: {url}\n{text[:2000]}\n"  # Limit to 2000 chars per source
                
        except Exception as e:
            print(f"Error extracting content from {url}: {e}")
            continue
    
    return content



@app.post("/analyze-document")
async def analyze_document(
    file: UploadFile = File(...),
    analysis_type: str = Form("all")
):
    """Analyze uploaded legal document"""
    try:
        content = await file.read()
        text = extract_text_from_content(content, file.filename)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="Could not extract text from the uploaded file.")
        
        results = {}
        
        if analysis_type in ["clauses", "all"]:
            results["clauses"] = extract_and_simplify_clauses(text)
        
        if analysis_type in ["entities", "all"]:
            results["entities"] = extract_entities(text)
        
        if analysis_type in ["classification", "all"]:
            results["document_type"] = classify_document(text)
        
        return results
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error analyzing document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error analyzing document: {str(e)}")

@app.post("/search-and-analyze")
async def search_and_analyze(
    question: str = Form(...),
    file: UploadFile = File(None),
    include_document_analysis: bool = Form(False)
):
    """Combined endpoint for search + document analysis"""
    try:
        results = {}
        
        # If file provided, analyze it and answer question based on document
        if file and include_document_analysis:
            content = await file.read()
            text = extract_text_from_content(content, file.filename)
            
            # Extract specific information from the document
            clauses = extract_clauses(text)
            entities = extract_entities(text)
            
            # Find relevant clauses and entities for the question
            relevant_clauses = []
            question_lower = question.lower()
            
            # Extract relevant clauses based on question keywords
            for clause in clauses:
                clause_text = clause["original"].lower()
                if any(keyword in clause_text for keyword in question_lower.split()):
                    relevant_clauses.append(clause)
            
            # Create a clean, professional document-based answer
            doc_answer = f"""**Document Analysis: {file.filename}**

**Question:** {question}

"""
            
            # Add relevant clauses with proper numbering
            if relevant_clauses:
                doc_answer += "**Relevant Clauses:**\n"
                for i, clause in enumerate(relevant_clauses[:3], 1):
                    doc_answer += f"{i}. **{clause['type'].title()}**: {clause['original']}\n"
                doc_answer += "\n"
            
            # Add key entities in a clean format
            if entities.get("parties"):
                doc_answer += "**Parties:** " + ", ".join(entities["parties"][:3]) + "\n\n"
            
            if entities.get("dates"):
                doc_answer += "**Key Dates:** " + ", ".join(entities["dates"][:3]) + "\n\n"
            
            if entities.get("amounts"):
                doc_answer += "**Financial Terms:** " + ", ".join(entities["amounts"][:3]) + "\n\n"
            
            # Add document summary
            doc_answer += f"**Document Type:** {classify_document(text)}\n"
            doc_answer += f"**Total Clauses:** {len(clauses)}\n\n"
            
            # Provide a concise answer
            if relevant_clauses:
                most_relevant = relevant_clauses[0]
                simplified = simplify_specific_clause(most_relevant['original'])
                doc_answer += f"**Answer:**\n{simplified}"
            else:
                doc_answer += "**Answer:** The document contains general legal terms. Please ask specific questions about particular clauses or terms."
            
            results["document_answer"] = doc_answer
            
            # Clean document analysis
            doc_analysis = {
                "clauses": extract_and_simplify_clauses(text),
                "entities": entities,
                "document_type": classify_document(text),
                "document_summary": f"Analyzed: {file.filename}",
                "total_clauses": len(clauses),
                "key_entities": {
                    "parties": len(entities.get("parties", [])),
                    "dates": len(entities.get("dates", [])),
                    "amounts": len(entities.get("amounts", []))
                }
            }
            
            results["document_analysis"] = doc_analysis
        else:
            # Fallback to general search
            search_results = await ask_question(Question(question=question))
            results["search_results"] = search_results
        
        return results
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error in search_and_analyze: {e}")
        raise HTTPException(status_code=500, detail=f"Error in combined analysis: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8001) 